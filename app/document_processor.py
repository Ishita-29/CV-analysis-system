import os
import io
import json
import base64
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
import tempfile

import fitz  # PyMuPDF
import docx
import pandas as pd
from PIL import Image, ImageFilter, ImageEnhance
import pytesseract
from pdf2image import convert_from_path

class DocumentProcessor:
    """
    Enhanced document processor with improved OCR capabilities for
    handling scanned PDFs and image-based documents.
    """
    
    def __init__(self, ocr_enabled: bool = True, ocr_lang: str = "eng"):
        """
        Initialize the document processor.
        
        Args:
            ocr_enabled: Whether to use OCR for image-based PDFs
            ocr_lang: Language for OCR (default: "eng")
        """
        self.ocr_enabled = ocr_enabled
        self.ocr_lang = ocr_lang
        self.supported_extensions = ['.pdf', '.docx', '.doc', '.jpg', '.jpeg', '.png']
        
        # Configure Tesseract path if environment variable is set
        tesseract_path = os.getenv("TESSERACT_PATH")
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
    def process_document(self, file_path: Union[str, Path]) -> dict:
        """
        Process a single CV document and extract raw text.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            dict: Document information including raw text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file format: {extension}")
        
        file_info = {
            "filename": file_path.name,
            "extension": extension,
            "file_size": file_path.stat().st_size,
            "path": str(file_path),
        }
        
        if extension == '.pdf':
            raw_text, is_scanned, images, page_count = self._process_pdf(file_path)
            file_info["page_count"] = page_count
        elif extension in ['.docx', '.doc']:
            raw_text = self._process_docx(file_path)
            is_scanned = False
            images = []
            file_info["page_count"] = None
        elif extension in ['.jpg', '.jpeg', '.png']:
            # Direct image processing
            raw_text, images = self._process_image(file_path)
            is_scanned = True
            file_info["page_count"] = 1
            
        file_info.update({
            "raw_text": raw_text,
            "is_scanned": is_scanned,
            "images": images,
            "ocr_applied": is_scanned and self.ocr_enabled,
            "text_quality": self._assess_text_quality(raw_text) if raw_text else "poor"
        })
        
        return file_info
    
    def _process_pdf(self, file_path: Path) -> Tuple[str, bool, List[dict], int]:
        """
        Process a PDF file with enhanced OCR capabilities.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            tuple: (raw_text, is_scanned, images, page_count)
        """
        pdf_document = fitz.open(file_path)
        raw_text = ""
        images = []
        text_extraction_confidence = 0
        
        # Check if PDF appears to be image-based
        is_likely_scanned = True
        text_content_pages = 0
        
        # First, try normal text extraction
        for page_num, page in enumerate(pdf_document):
            page_text = page.get_text()
            if len(page_text.strip()) > 100:
                text_content_pages += 1
                text_extraction_confidence += 1
            raw_text += page_text + "\n\n"
        
        # If enough pages have text content, it's probably not a scanned PDF
        if text_content_pages / len(pdf_document) > 0.3:
            is_likely_scanned = False
        
        # If it's likely a scanned PDF and OCR is enabled, apply OCR
        if is_likely_scanned and self.ocr_enabled:
            ocr_text = ""
            
            # Use pdf2image to convert PDF to images for better OCR
            with tempfile.TemporaryDirectory() as temp_dir:
                try:
                    pdf_images = convert_from_path(
                        file_path,
                        dpi=300,
                        output_folder=temp_dir,
                        fmt='jpeg',
                        thread_count=os.cpu_count() or 1
                    )
                    
                    for page_num, img in enumerate(pdf_images):
                        # Enhance image for better OCR
                        enhanced_img = self._enhance_image_for_ocr(img)
                        
                        # Perform OCR
                        page_ocr_text = pytesseract.image_to_string(
                            enhanced_img,
                            lang=self.ocr_lang,
                            config='--psm 1 --oem 3'  # Automatic page segmentation with OEM
                        )
                        
                        ocr_text += page_ocr_text + "\n\n"
                        
                        # Save image for reference
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='JPEG', quality=50)  # Lower quality to save space
                        img_base64 = base64.b64encode(img_byte_arr.getvalue()).decode('utf-8')
                        
                        images.append({
                            "page": page_num + 1,
                            "mime_type": "image/jpeg",
                            "data": img_base64
                        })
                    
                    # If OCR text has more content, use it instead
                    if len(ocr_text.strip()) > len(raw_text.strip()) * 1.2:
                        raw_text = ocr_text
                except Exception as e:
                    print(f"Error during OCR processing: {str(e)}")
                    # If OCR fails, we'll use the text we already extracted
                    pass
                    
        is_scanned = is_likely_scanned
        
        return raw_text, is_scanned, images, len(pdf_document)
    
    def _enhance_image_for_ocr(self, img: Image) -> Image:
        """
        Enhance image quality for better OCR results.
        
        Args:
            img: PIL Image
        
        Returns:
            PIL Image: Enhanced image
        """
        # Convert to grayscale
        img = img.convert('L')
        
        # Increase contrast
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)
        
        # Apply slight blur to reduce noise
        img = img.filter(ImageFilter.GaussianBlur(radius=0.5))
        
        # Apply threshold to make text more distinct
        # Use a local adaptive threshold
        def local_threshold(img, block_size=35, c=10):
            # Simple adaptive thresholding
            img_arr = np.array(img)
            height, width = img_arr.shape
            output = np.zeros_like(img_arr)
            
            for i in range(height):
                for j in range(width):
                    # Get block boundaries
                    block_start_i = max(0, i - block_size//2)
                    block_end_i = min(height, i + block_size//2 + 1)
                    block_start_j = max(0, j - block_size//2)
                    block_end_j = min(width, j + block_size//2 + 1)
                    
                    # Calculate local threshold
                    block = img_arr[block_start_i:block_end_i, block_start_j:block_end_j]
                    threshold = np.mean(block) - c
                    
                    # Apply threshold
                    output[i, j] = 0 if img_arr[i, j] < threshold else 255
            
            return Image.fromarray(output)
        
        try:
            import numpy as np
            img = local_threshold(img)
        except ImportError:
            # Fallback to simple thresholding if numpy isn't available
            img = img.point(lambda x: 0 if x < 128 else 255, '1')
        
        return img
    
    def _process_docx(self, file_path: Path) -> str:
        """
        Process a DOCX file and extract text.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            str: Raw text from the document
        """
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        
        # Also extract tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                tables_text.append(" | ".join(row_text))
        
        # Combine all text
        all_text = "\n\n".join(paragraphs)
        
        if tables_text:
            all_text += "\n\nTABLE DATA:\n" + "\n".join(tables_text)
            
        return all_text
    
    def _process_image(self, file_path: Path) -> Tuple[str, List[dict]]:
        """
        Process an image file with OCR.
        
        Args:
            file_path: Path to the image file
            
        Returns:
            tuple: (raw_text, images)
        """
        if not self.ocr_enabled:
            return "", []
            
        images = []
        raw_text = ""
        
        try:
            # Open and enhance image
            img = Image.open(file_path)
            enhanced_img = self._enhance_image_for_ocr(img)
            
            # Perform OCR
            raw_text = pytesseract.image_to_string(
                enhanced_img,
                lang=self.ocr_lang,
                config='--psm 1 --oem 3'
            )
            
            # Save image for reference
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='JPEG', quality=50)
            img_base64 = base64.b64encode(img_byte_arr.getvalue()).decode('utf-8')
            
            images.append({
                "page": 1,
                "mime_type": f"image/{img.format.lower() if img.format else 'jpeg'}",
                "data": img_base64
            })
            
        except Exception as e:
            print(f"Error processing image: {str(e)}")
            
        return raw_text, images
    
    def _assess_text_quality(self, text: str) -> str:
        """
        Assess the quality of extracted text.
        
        Args:
            text: Extracted text
            
        Returns:
            str: Quality assessment ('excellent', 'good', 'fair', 'poor')
        """
        # Check if text is empty
        if not text or len(text.strip()) < 50:
            return "poor"
        
        # Check for common OCR errors and gibberish
        gibberish_indicators = ['€', '™', '�', '□', '■', '≤', '≥', '¢', '±', '∞']
        gibberish_count = sum(text.count(char) for char in gibberish_indicators)
        
        # Calculate text-to-special-char ratio
        special_chars = sum(not c.isalnum() and not c.isspace() for c in text)
        text_length = len(text)
        spec_char_ratio = special_chars / text_length if text_length > 0 else 1
        
        # Check for words without vowels (often OCR errors)
        words = text.split()
        words_without_vowels = 0
        for word in words:
            if len(word) > 3 and not any(vowel in word.lower() for vowel in 'aeiou'):
                words_without_vowels += 1
        
        no_vowel_ratio = words_without_vowels / len(words) if words else 1
        
        # Make assessment
        if gibberish_count < 5 and spec_char_ratio < 0.1 and no_vowel_ratio < 0.05:
            return "excellent"
        elif gibberish_count < 15 and spec_char_ratio < 0.15 and no_vowel_ratio < 0.1:
            return "good"
        elif gibberish_count < 30 and spec_char_ratio < 0.2 and no_vowel_ratio < 0.15:
            return "fair"
        else:
            return "poor"
    
    def batch_process(self, directory: Union[str, Path]) -> List[dict]:
        """
        Process all CV documents in a directory.
        
        Args:
            directory: Path to the directory containing CV documents
            
        Returns:
            list: List of documents information dictionaries
        """
        directory = Path(directory)
        
        if not directory.exists() or not directory.is_dir():
            raise NotADirectoryError(f"Directory not found: {directory}")
            
        documents = []
        
        for file_path in directory.iterdir():
            if file_path.suffix.lower() in self.supported_extensions:
                try:
                    doc_info = self.process_document(file_path)
                    documents.append(doc_info)
                except Exception as e:
                    print(f"Error processing {file_path}: {str(e)}")
                    
        return documents
    
    def save_processed_documents(self, documents: List[dict], output_dir: Union[str, Path]):
        """
        Save the processed documents to the output directory.
        
        Args:
            documents: List of processed document information
            output_dir: Directory to save the processed data
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        for doc in documents:
            # Save document info without large binary data
            doc_info = doc.copy()
            # Remove base64 image data to keep files small
            if 'images' in doc_info:
                doc_info['image_count'] = len(doc_info['images'])
                doc_info.pop('images')
                
            output_file = output_dir / f"{Path(doc['filename']).stem}.json"
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(doc_info, f, ensure_ascii=False, indent=2)
                
        # Save a summary CSV
        summary_data = [{
            'filename': doc['filename'],
            'extension': doc['extension'],
            'file_size': doc['file_size'],
            'is_scanned': doc.get('is_scanned', False),
            'ocr_applied': doc.get('ocr_applied', False),
            'text_quality': doc.get('text_quality', 'unknown'),
            'text_length': len(doc.get('raw_text', '')),
            'page_count': doc.get('page_count', None)
        } for doc in documents]
        
        summary_df = pd.DataFrame(summary_data)
        summary_df.to_csv(output_dir / 'document_summary.csv', index=False)
        
    def export_to_excel(self, parsed_cvs: Dict[str, Dict], output_file: Union[str, Path]):
        """
        Export parsed CV data to Excel with multiple sheets.
        
        Args:
            parsed_cvs: Dictionary mapping CV IDs to parsed CV data
            output_file: Path to the output Excel file
        """
        output_file = Path(output_file)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Create a Pandas Excel writer
        writer = pd.ExcelWriter(output_file, engine='xlsxwriter')
        
        # Create summary sheet
        summary_data = []
        for cv_id, cv_data in parsed_cvs.items():
            personal_info = cv_data.get('personal_info', {})
            education = cv_data.get('education', [])
            experience = cv_data.get('experience', [])
            skills = cv_data.get('skills', [])
            
            highest_degree = None
            if education:
                # Simple heuristic to find highest degree
                degree_texts = [edu.get('degree', '') for edu in education]
                for degree in ['phd', 'doctor', 'master', 'bachelor']:
                    for text in degree_texts:
                        if degree in text.lower():
                            highest_degree = text
                            break
                    if highest_degree:
                        break
            
            # Calculate total experience years (rough estimate)
            total_exp_years = 0
            for exp in experience:
                date_range = exp.get('date_range', '')
                if '-' in date_range:
                    years = date_range.split('-')
                    if len(years) == 2:
                        try:
                            start_year = int(''.join(filter(str.isdigit, years[0])))
                            end_year = 2023 if 'present' in years[1].lower() else int(''.join(filter(str.isdigit, years[1])))
                            total_exp_years += (end_year - start_year)
                        except ValueError:
                            pass
            
            summary_data.append({
                'Name': personal_info.get('name', 'Unknown'),
                'Email': personal_info.get('email', ''),
                'Phone': personal_info.get('phone', ''),
                'Location': personal_info.get('location', ''),
                'Highest Degree': highest_degree or 'Not specified',
                'Years of Experience': total_exp_years,
                'Top Skills': ', '.join(skills[:5]) if skills else '',
                'CV ID': cv_id
            })
        
        # Convert to DataFrame and save to Excel
        summary_df = pd.DataFrame(summary_data)
        summary_df.to_excel(writer, sheet_name='Summary', index=False)
        
        # Create sheets for other categories
        
        # Education Sheet
        education_data = []
        for cv_id, cv_data in parsed_cvs.items():
            name = cv_data.get('personal_info', {}).get('name', 'Unknown')
            for edu in cv_data.get('education', []):
                education_data.append({
                    'Name': name,
                    'CV ID': cv_id,
                    'Degree': edu.get('degree', ''),
                    'Institution': edu.get('institution', ''),
                    'Date Range': edu.get('date_range', '')
                })
        
        if education_data:
            pd.DataFrame(education_data).to_excel(writer, sheet_name='Education', index=False)
        
        # Experience Sheet
        experience_data = []
        for cv_id, cv_data in parsed_cvs.items():
            name = cv_data.get('personal_info', {}).get('name', 'Unknown')
            for exp in cv_data.get('experience', []):
                experience_data.append({
                    'Name': name,
                    'CV ID': cv_id,
                    'Job Title': exp.get('job_title', ''),
                    'Company': exp.get('company', ''),
                    'Date Range': exp.get('date_range', '')
                })
        
        if experience_data:
            pd.DataFrame(experience_data).to_excel(writer, sheet_name='Experience', index=False)
        
        # Skills Sheet
        skills_data = []
        for cv_id, cv_data in parsed_cvs.items():
            name = cv_data.get('personal_info', {}).get('name', 'Unknown')
            skills = cv_data.get('skills', [])
            
            if skills:
                for skill in skills:
                    skills_data.append({
                        'Name': name,
                        'CV ID': cv_id,
                        'Skill': skill
                    })
        
        if skills_data:
            pd.DataFrame(skills_data).to_excel(writer, sheet_name='Skills', index=False)
        
        # Save and close the Excel file
        writer.close()
        
        return output_file